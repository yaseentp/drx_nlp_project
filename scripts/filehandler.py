from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import fitz
from collections import defaultdict
from pathlib import Path
from typing import List, Dict, Any
import uuid


def detect_file_type_and_extract_text(file_path):
    ext = Path(file_path).suffix.lower()
    if ext == '.docx':
        return extract_docx_structure(file_path)
    elif ext == '.pdf':
        return extract_pdf_structured_json(file_path)
    # elif ext in ['.png', '.jpg', '.jpeg']:
    #     return extract_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

# def iter_block_items(parent):
#     """Yield paragraphs and tables in document order."""
#     for child in parent.element.body.iterchildren():
#         if child.tag.endswith('p'):
#             yield Paragraph(child, parent)
#         elif child.tag.endswith('tbl'):
#             yield Table(child, parent)



from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from typing import List, Dict, Any
import uuid

def extract_docx_structure(filepath: str) -> List[Dict[str, Any]]:
    document = Document(filepath)
    hierarchy = []
    paragraphs_per_page = 20  # heuristic

    def detect_heading_level(style_name: str) -> int:
        if style_name.startswith("Heading"):
            try:
                return int(style_name.split(" ")[1])
            except:
                return 0
        return 0

    def is_list_style(style_name: str) -> bool:
        return any(kw in style_name.lower() for kw in ['list', 'bullet', 'number'])

    def get_list_indent_level(para) -> int:
        try:
            return para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else 0
        except:
            return 0

    def get_formatting(para) -> Dict[str, bool]:
        formatting = {"bold": False, "italic": False, "underline": False}
        for run in para.runs:
            formatting["bold"] |= run.bold or False
            formatting["italic"] |= run.italic or False
            formatting["underline"] |= run.underline or False
        return formatting

    def build_structured_obj(
        obj_type: str,
        text: str,
        section_index: int,
        page_number: int,
        style: str = None,
        heading_level: int = None,
        formatting: Dict[str, bool] = None,
        list_level: int = None
    ) -> Dict[str, Any]:
        return {
            "id": str(uuid.uuid4()),
            "type": obj_type,
            "text": text,
            "style": style,
            "heading_level": heading_level,
            "formatting": formatting or {},
            "list_level": list_level,
            "position": {
                "section_index": section_index,
                "page_number": page_number
            },
            "children": []
        }

    def iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if child.tag.endswith('tbl'):
                yield {'type': 'table', 'item': Table(child, parent)}
            elif child.tag.endswith('p'):
                yield {'type': 'paragraph', 'item': Paragraph(child, parent)}

    def parse_table(table, section_index, page_number):
        table_data = []
        for i, row in enumerate(table.rows):
            row_data = []
            for j, cell in enumerate(row.cells):
                row_data.append({
                    'text': cell.text.strip(),
                    'row': i,
                    'col': j
                })
            table_data.append(row_data)

        return {
            'id': str(uuid.uuid4()),
            'type': 'Table',
            'position': {
                'section_index': section_index,
                'page_number': page_number
            },
            'content': table_data,
            'children': []
        }

    def parse_paragraph(para, section_index, page_number):
        text = para.text.strip()
        if not text:
            return None

        style = para.style.name if para.style else ''
        heading_level = detect_heading_level(style)
        para_type = 'Paragraph'
        if heading_level:
            para_type = 'Heading'
        elif is_list_style(style):
            para_type = 'List'

        return build_structured_obj(
            obj_type=para_type,
            text=text,
            style=style,
            heading_level=heading_level,
            formatting=get_formatting(para),
            section_index=section_index,
            page_number=page_number,
            list_level=get_list_indent_level(para)
        )

    def extract_headers_and_footers():
        for i, section in enumerate(document.sections):
            if section.header.is_linked_to_previous and i > 0:
                continue
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        hierarchy.append(build_structured_obj(
                            obj_type="Header",
                            text=text,
                            section_index=i,
                            page_number=1
                        ))
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        hierarchy.append(build_structured_obj(
                            obj_type="Footer",
                            text=text,
                            section_index=i,
                            page_number=-1
                        ))

    # --- main execution ---
    extract_headers_and_footers()
    section_index = 0
    paragraph_counter = 0

    for block in iter_block_items(document):
        estimated_page = paragraph_counter // paragraphs_per_page + 1

        if block['type'] == 'paragraph':
            parsed = parse_paragraph(block['item'], section_index, estimated_page)
            if parsed:
                hierarchy.append(parsed)
                paragraph_counter += 1

        elif block['type'] == 'table':
            hierarchy.append(parse_table(block['item'], section_index, estimated_page))

        section_index += 1

    return hierarchy


def extract_pdf_structured_json(pdf_path):
    doc = fitz.open(pdf_path)
    results = []
    section_index = 0

    top_texts = defaultdict(int)
    bottom_texts = defaultdict(int)

    # Pass 1: detect common headers/footers
    for page in doc:
        height = page.rect.height
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                y = line["bbox"][1]
                text = "".join(span["text"] for span in line["spans"]).strip()
                if not text:
                    continue
                if y < height * 0.1:
                    top_texts[text] += 1
                elif y > height * 0.9:
                    bottom_texts[text] += 1

    common_top = {k for k, v in top_texts.items() if v >= len(doc) * 0.8}
    common_bottom = {k for k, v in bottom_texts.items() if v >= len(doc) * 0.8}

    for page_number, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        lines_info = []

        # Table detection
        for table in page.find_tables():
            matrix = table.extract()
            if not matrix:
                continue
            table_text = "\n".join(["\t".join(cell if cell is not None else "" for cell in row) for row in matrix])
            results.append({
                "id": str(uuid.uuid4()),
                "type": "Table",
                "text": table_text,
                "style": "Table Grid",
                "heading_level": 0,
                "formatting": {"bold": False, "italic": False, "underline": False},
                "list_level": 0,
                "position": {"section_index": section_index, "page_number": page_number},
                "children": []
            })

        # Process non-table text
        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                y = line["bbox"][1]
                text = "".join(span["text"] for span in line["spans"]).strip()
                if not text or text in common_top or text in common_bottom:
                    continue
                fonts = [span["font"] for span in line["spans"]]
                lines_info.append({"text": text, "y": y, "fonts": fonts})

        lines_info.sort(key=lambda x: x["y"])
        prev_y = None
        paragraph_buffer = []

        for line in lines_info:
            text = line["text"]
            fonts = line["fonts"]
            y = line["y"]

            bold_count = sum("bold" in f.lower() or "bd" in f.lower() for f in fonts)
            is_header = bold_count / len(fonts) > 0.5 if fonts else False

            if is_header:
                results.append({
                    "id": str(uuid.uuid4()),
                    "type": "Heading",
                    "text": text,
                    "style": "Heading 1",
                    "heading_level": 1,
                    "formatting": {"bold": True, "italic": False, "underline": False},
                    "list_level": 0,
                    "position": {"section_index": section_index, "page_number": page_number},
                    "children": []
                })
                section_index += 1
                continue

            if prev_y is not None and abs(y - prev_y) > 10 and paragraph_buffer:
                results.append({
                    "id": str(uuid.uuid4()),
                    "type": "Paragraph",
                    "text": " ".join(paragraph_buffer),
                    "style": "Normal",
                    "heading_level": 0,
                    "formatting": {"bold": False, "italic": False, "underline": False},
                    "list_level": 0,
                    "position": {"section_index": section_index, "page_number": page_number},
                    "children": []
                })
                paragraph_buffer = []

            paragraph_buffer.append(text)
            prev_y = y

            if text[-1] in ".?!":
                results.append({
                    "id": str(uuid.uuid4()),
                    "type": "Paragraph",
                    "text": " ".join(paragraph_buffer),
                    "style": "Normal",
                    "heading_level": 0,
                    "formatting": {"bold": False, "italic": False, "underline": False},
                    "list_level": 0,
                    "position": {"section_index": section_index, "page_number": page_number},
                    "children": []
                })
                paragraph_buffer = []

        if paragraph_buffer:
            results.append({
                "id": str(uuid.uuid4()),
                "type": "Paragraph",
                "text": " ".join(paragraph_buffer),
                "style": "Normal",
                "heading_level": 0,
                "formatting": {"bold": False, "italic": False, "underline": False},
                "list_level": 0,
                "position": {"section_index": section_index, "page_number": page_number},
                "children": []
            })

    return results


